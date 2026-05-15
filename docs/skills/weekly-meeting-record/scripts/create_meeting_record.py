import argparse
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


WEEKDAY_MAP = {
    0: "星期一",
    1: "星期二",
    2: "星期三",
    3: "星期四",
    4: "星期五",
    5: "星期六",
    6: "星期日",
}


def load_payload(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def format_people(values) -> str:
    if not values:
        return "无"
    if isinstance(values, str):
        return values.strip() or "无"
    cleaned = [str(v).strip() for v in values if str(v).strip()]
    if not cleaned:
        return "无"
    if len(cleaned) == 1 and cleaned[0] == "无":
        return "无"
    return "、".join(cleaned)


def normalize_terminal_punctuation(text: str, is_last: bool) -> str:
    cleaned = text.strip().rstrip("；;。.")
    return f"{cleaned}{'。' if is_last else '；'}"


def set_cell_border(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        elem = tc_borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            tc_borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")


def set_run_font(run, east_asia: str, size: int, bold: bool = False) -> None:
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(cell, text: str = "", *, bold: bool = False, align: int = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph = cell.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 12, bold=bold)


def add_time_paragraph(cell, label: str, value: str) -> None:
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    label_run = paragraph.add_run(label)
    set_run_font(label_run, "方正小标宋简体", 12)

    colon_run = paragraph.add_run("：")
    set_run_font(colon_run, "方正小标宋简体", 12, bold=True)

    value_run = paragraph.add_run(value)
    set_run_font(value_run, "方正小标宋简体", 12)


def clear_default_paragraph(cell) -> None:
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]._element
        p.getparent().remove(p)


def derive_weekday(date_text: str) -> str:
    match = re.match(r"^\s*(\d{4})年(\d{1,2})月(\d{1,2})日\s*$", date_text)
    if not match:
        return ""
    year, month, day = map(int, match.groups())
    return WEEKDAY_MAP[datetime(year, month, day).weekday()]


def build_time_display(data: dict) -> str:
    weekday = data.get("weekday") or derive_weekday(data["date"])
    if weekday:
        return f"{data['date']}{weekday} {data['time']}"
    return f"{data['date']} {data['time']}"


def write_meta_row(cells, label_text: str, value_text: str) -> None:
    label_cell, value_cell = cells
    clear_default_paragraph(label_cell)
    clear_default_paragraph(value_cell)
    label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    label_paragraph = label_cell.add_paragraph()
    label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label_paragraph.add_run(label_text)
    set_run_font(label_run, "方正小标宋简体", 12)

    value_paragraph = value_cell.add_paragraph()
    value_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    value_run = value_paragraph.add_run(value_text)
    set_run_font(value_run, "宋体", 12)


def build_document(data: dict) -> Document:
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(12)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.175)
    section.right_margin = Cm(3.48)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run(data["title"])
    set_run_font(title_run, "方正小标宋简体", 18, bold=True)

    table = doc.add_table(rows=8, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(2.4), Cm(7.2), Cm(2.4), Cm(5.8)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            set_cell_border(row.cells[idx])

    time_row = table.rows[0].cells
    write_meta_row((time_row[0], time_row[1]), "时间", build_time_display(data))
    write_meta_row((time_row[2], time_row[3]), "地点", data["location"])

    leave_row = table.rows[1].cells
    leave_merged = leave_row[1].merge(leave_row[3])
    write_meta_row((leave_row[0], leave_merged), "请假人员", format_people(data.get("leave_people")))

    late_row = table.rows[2].cells
    late_merged = late_row[1].merge(late_row[3])
    write_meta_row((late_row[0], late_merged), "迟到人员", format_people(data.get("late_people")))

    absence_row = table.rows[3].cells
    absence_merged = absence_row[1].merge(absence_row[3])
    write_meta_row((absence_row[0], absence_merged), "旷会人员", format_people(data.get("absence_people")))

    host_row = table.rows[4].cells
    write_meta_row((host_row[0], host_row[1]), "主持人", format_people(data.get("hosts")))
    write_meta_row((host_row[2], host_row[3]), "记录人", str(data["recorder"]).strip())

    heading_row = table.rows[5].cells
    heading_cell = heading_row[0].merge(heading_row[3])
    clear_default_paragraph(heading_cell)
    heading_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    heading_paragraph = heading_cell.add_paragraph()
    heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading_paragraph.add_run("会议内容")
    set_run_font(heading_run, "方正小标宋简体", 12)

    content_row = table.rows[6].cells
    content_cell = content_row[0].merge(content_row[3])
    clear_default_paragraph(content_cell)

    add_paragraph(content_cell, "")
    add_time_paragraph(content_cell, "开始时间", data["start_time"])
    add_paragraph(content_cell, "")

    sections = data.get("sections", [])
    for section_index, section_data in enumerate(sections, start=1):
        add_paragraph(content_cell, f"{section_index}.{section_data['title']}：", bold=True)
        items = section_data.get("items", [])
        for item_index, item in enumerate(items, start=1):
            add_paragraph(content_cell, normalize_terminal_punctuation(item, item_index == len(items)))
        if section_index != len(sections):
            add_paragraph(content_cell, "")

    add_paragraph(content_cell, "")
    add_time_paragraph(content_cell, "结束时间", data["end_time"])
    add_paragraph(content_cell, "")

    footer_row = table.rows[7].cells
    footer_cell = footer_row[0].merge(footer_row[3])
    clear_default_paragraph(footer_cell)
    add_paragraph(footer_cell, "")

    return doc


def main() -> None:
    parser = argparse.ArgumentParser(description="Create a meeting-record docx from JSON.")
    parser.add_argument("--input", required=True, help="Path to the JSON payload.")
    parser.add_argument("--output", required=True, help="Path to the output docx file.")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    payload = load_payload(input_path)
    doc = build_document(payload)
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
